import logging
import os
import pandas as pd
from pathlib import Path
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


def ensure_directories():
    """Ensure output and upload directories exist."""
    os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)


def parse_input_file(file_path: str) -> List[Dict[str, str]]:
    """Parse CSV or Excel file and return list of book entries."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".csv":
            df = pd.read_csv(file_path)
        elif suffix in (".xlsx", ".xls"):
            df = pd.read_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Use .csv, .xlsx, or .xls")

        # Normalize column names
        df.columns = [col.strip().lower().replace(" ", "_") for col in df.columns]

        required_columns = ["title"]
        for col in required_columns:
            if col not in df.columns:
                raise ValueError(f"Missing required column: '{col}'. Found columns: {list(df.columns)}")

        books = []
        for _, row in df.iterrows():
            book_entry = {
                "title": str(row.get("title", "")).strip(),
                "notes_on_outline_before": str(row.get("notes_on_outline_before", "")).strip()
                if pd.notna(row.get("notes_on_outline_before"))
                else None,
            }
            if book_entry["title"]:
                books.append(book_entry)

        logger.info(f"Parsed {len(books)} books from {file_path}")
        return books

    except Exception as e:
        logger.error(f"Failed to parse file {file_path}: {e}")
        raise


def export_book_txt(book_id: str, title: str, chapters: list) -> str:
    """Export book as a .txt file. Returns the file path."""
    ensure_directories()
    file_path = os.path.join(settings.OUTPUT_DIR, f"book_{book_id}.txt")

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(f"{'=' * 60}\n")
        f.write(f"{title.upper()}\n")
        f.write(f"{'=' * 60}\n\n")

        for chapter in chapters:
            f.write(f"\n{'─' * 40}\n")
            ch_title = chapter.get("chapter_title", f"Chapter {chapter['chapter_number']}")
            f.write(f"Chapter {chapter['chapter_number']}: {ch_title}\n")
            f.write(f"{'─' * 40}\n\n")
            f.write(chapter.get("content", "") + "\n")

        f.write(f"\n{'=' * 60}\n")
        f.write("END OF BOOK\n")
        f.write(f"{'=' * 60}\n")

    logger.info(f"Exported TXT: {file_path}")
    return file_path


def export_book_docx(book_id: str, title: str, chapters: list) -> str:
    """Export book as a .docx file. Returns the file path."""
    ensure_directories()
    file_path = os.path.join(settings.OUTPUT_DIR, f"book_{book_id}.docx")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Title page
    doc.add_paragraph("")
    doc.add_paragraph("")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(28)
    run.bold = True

    doc.add_page_break()

    # Chapters
    for chapter in chapters:
        ch_title = chapter.get("chapter_title", f"Chapter {chapter['chapter_number']}")
        heading = doc.add_heading(f"Chapter {chapter['chapter_number']}: {ch_title}", level=1)

        content = chapter.get("content", "")
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    doc.save(file_path)
    logger.info(f"Exported DOCX: {file_path}")
    return file_path


def parse_outline_to_chapters(outline: str) -> List[Dict[str, str]]:
    """Parse the generated outline into a list of chapter dicts."""
    chapters = []
    lines = outline.strip().split("\n")
    current_chapter = None
    current_description = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Strip markdown bold markers (**) and other formatting
        clean_line = line.strip("*").strip("#").strip().strip("*").strip()

        # Match patterns like "Chapter 1:", "Chapter 1.", "1.", "1:"
        is_chapter_line = False
        chapter_number = None
        chapter_title = ""

        if clean_line.lower().startswith("chapter"):
            parts = clean_line.split(":", 1)
            if len(parts) >= 1:
                num_part = parts[0].lower().replace("chapter", "").strip().rstrip(".")
                try:
                    chapter_number = int(num_part)
                    chapter_title = parts[1].strip().strip("*").strip() if len(parts) > 1 else ""
                    is_chapter_line = True
                except ValueError:
                    pass

        if is_chapter_line and chapter_number is not None:
            if current_chapter is not None:
                current_chapter["description"] = " ".join(current_description).strip()
                chapters.append(current_chapter)
            current_chapter = {
                "chapter_number": chapter_number,
                "chapter_title": chapter_title,
                "description": "",
            }
            current_description = []
        elif current_chapter is not None:
            current_description.append(line)

    # Add last chapter
    if current_chapter is not None:
        current_chapter["description"] = " ".join(current_description).strip()
        chapters.append(current_chapter)

    logger.info(f"Parsed {len(chapters)} chapters from outline")
    return chapters
